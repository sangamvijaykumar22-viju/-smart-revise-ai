from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "single", "color": "#0000FF"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def create_report():
    doc = Document()

    # Title Page
    section = doc.sections[0]
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run('PROJECT REPORT ON')
    run.font.size = Pt(16)
    run.bold = True

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run('\nSMART REVISE: AI-POWERED LEARNING SYSTEM\nWITH HARDWARE INTEGRATION')
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph('\n' * 8)
    
    tech_para = doc.add_paragraph()
    tech_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tech_para.add_run('Tech Stack: Python, Flask, SQLite, AI (LLMs), Raspberry Pi')
    run.font.size = Pt(12)
    run.italic = True

    doc.add_page_break()

    # Table of Contents
    doc.add_heading('TABLE OF CONTENTS', level=1)
    toc_data = [
        ("CHAPTER-1: INTRODUCTION", "1"),
        ("1.1 Purpose of the project", "2"),
        ("1.2 Problem with Existing Systems", "2"),
        ("1.3 Proposed System", "3"),
        ("1.4 Scope of the Project", "3"),
        ("1.5 Architecture Diagram", "4"),
        ("CHAPTER–2: LITERATURE SURVEY", "5"),
        ("CHAPTER-3: SOFTWARE REQUIREMENT SPECIFICATION", "9"),
        ("3.4 Functional Requirements", "10"),
        ("3.7 Software Requirements", "11"),
        ("3.8 Hardware Requirements", "12"),
        ("CHAPTER–4: SYSTEM DESIGN", "14"),
        ("4.3 TECHNOLOGIES USED", "21"),
        ("CHAPTER–5: IMPLEMENTATION", "33"),
        ("5.1 Raspberry Pi Setup", "33"),
        ("5.2 Coding Logic", "35"),
        ("CHAPTER–6: SOFTWARE TESTING", "45"),
        ("CONCLUSION", "52"),
        ("BIBLIOGRAPHY", "55")
    ]
    
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    for item, pg in toc_data:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = pg

    doc.add_page_break()

    # Chapter 1
    doc.add_heading('CHAPTER-1: INTRODUCTION', level=1)
    doc.add_heading('1.1 Purpose of the project', level=2)
    doc.add_paragraph("The purpose of Smart Revise is to provide students with a personalized learning assistant that uses AI to generate content and IoT to maintain study discipline.")
    
    doc.add_heading('1.2 Problem with Existing Systems', level=2)
    doc.add_paragraph("Traditional learning methods are static and lack real-time engagement. Digital tools are often fragmented and don't provide physical study cues.")
    
    doc.add_heading('1.3 Proposed System', level=2)
    doc.add_paragraph("Smart Revise integrates an AI-powered chatbot, an automated note generator, and a hardware-integrated study planner to create a holistic educational environment.")

    doc.add_heading('1.5 Architecture Diagram', level=2)
    doc.add_paragraph("The system follows a layered architecture as described below:")
    
    # Architecture Table (Mimicking the image structure)
    arch_table = doc.add_table(rows=5, cols=2)
    arch_table.style = 'Table Grid'
    
    def fill_cell(row, col, title, content):
        cell = arch_table.rows[row].cells[col]
        p = cell.paragraphs[0]
        r = p.add_run(f"{title}\n")
        r.bold = True
        cell.add_paragraph(content)

    fill_cell(0, 0, "FRONTEND LAYER", "HTML5, CSS3, JavaScript (Vanilla)\n- UI + User Interaction")
    fill_cell(0, 1, "APPLICATION LAYER", "Python Flask Backend\n- RESTful APIs\n- Routing & Logic")
    fill_cell(1, 0, "DATABASE LAYER", "SQLite (SQLAlchemy ORM)\n- Users, Tasks, Revisions Data")
    fill_cell(1, 1, "AUTH SERVICE", "Flask-Bcrypt (Security)\n- JWT Session Management")
    fill_cell(2, 0, "AI ENGINE LAYER", "Large Language Models (LLM)\n- Note Generation\n- Chatbot Assistant")
    fill_cell(2, 1, "HARDWARE LAYER", "Raspberry Pi 4\n- GPIO Notification System\n- LED Study Reminders")
    
    doc.add_page_break()

    # Chapter 4: Technologies Used
    doc.add_heading('CHAPTER–4: SYSTEM DESIGN', level=1)
    doc.add_heading('4.3 TECHNOLOGIES USED', level=2)
    
    techs = [
        ("Python (Flask)", "Backend framework for handling API requests and business logic."),
        ("SQLite", "Lightweight relational database for local storage of student data."),
        ("SQLAlchemy", "ORM for seamless database interactions."),
        ("LLM (Groq/OpenAI)", "Advanced AI models used for generating custom revision notes."),
        ("Raspberry Pi", "Microcontroller for physical study-habit indicators."),
        ("Bcrypt", "Industry-standard password hashing for user security.")
    ]
    
    for t, desc in techs:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(f"{t}: ")
        r.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # Implementation
    doc.add_heading('CHAPTER–5: IMPLEMENTATION', level=1)
    doc.add_heading('5.2 Coding the logic', level=2)
    doc.add_paragraph("The backend implementation uses Flask Blueprints to separate Authentication, Revision Generation, and Planner logic.")
    
    doc.add_heading('5.3 Global Database Caching (Wiki-Model)', level=2)
    doc.add_paragraph("To manage API token limits, the system implements a Global Response Cache. If a student searches for a topic already generated by another user, the system retrieves it from the local database instantly, reducing token costs by up to 90% as the platform grows.")

    # Save
    doc.save('Smart_Revise_Report.docx')
    print("Report generated successfully: Smart_Revise_Report.docx")

if __name__ == "__main__":
    create_report()
